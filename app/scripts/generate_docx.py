import os
import sys

# Safe project-root bootstrap for direct execution
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

def main():
    print("==========================================")
    print("MedMitra Technical Documentation Compiler")
    print("==========================================\n")

    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        print("Note: 'python-docx' is not installed in the current environment.")
        print("To generate the Word (.docx) technical documentation file, run:")
        print("pip install python-docx")
        print("python -m scripts.generate_docx")
        sys.exit(0)

    docs_dir = os.path.join(project_root, "docs")
    os.makedirs(docs_dir, exist_ok=True)
    output_path = os.path.join(docs_dir, "COMPLETE_TECHNICAL_DOCUMENTATION.docx")

    print("Compiling COMPLETE_TECHNICAL_DOCUMENTATION.docx...")

    doc = Document()
    
    # Define styles and properties
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('MedMitra ML/Search Service', level=0)
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    # Subtitle
    subtitle = doc.add_paragraph('Complete Technical Documentation & System Reference Guide')
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph(f"Generated Date: {os.popen('date /t').read().strip() or '2026-06-03'}")
    doc.add_page_break()

    # Section 1: Executive Summary
    h1 = doc.add_heading('1. Executive Summary & Purpose', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph(
        "MedMitra ML Search Service is a production-grade backend microservice designed for healthcare symptom "
        "matching, clinical guidelines retrieval, deterministic emergency filtering, and prescription OCR text preparation."
    )
    doc.add_paragraph(
        "It provides safe, symptom-based educational context based on officially verified guidelines. It does NOT "
        "diagnose, prescribe drugs, recommend dose alterations, or replace clinical consultations. It operates "
        "entirely state-free and database-agnostic."
    )

    # Section 2: Clinical Safety Boundaries
    h2 = doc.add_heading('2. Clinical Safety Boundaries & Guardrails', level=1)
    h2.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph("Strict healthcare protocols are implemented at all software layers:")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Infant Safety Guard: ").bold = True
    p.add_run("Queries concerning patients under 2 months of age trigger immediate pediatrician referral errors.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Emergency Proximity Filter: ").bold = True
    p.add_run("High-risk symptoms bypass Qdrant and LLMs completely, returning immediate escalation messages.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("No Diagnostic Certainty: ").bold = True
    p.add_run("Retrievals expose similarity as 'retrieval_relevance' (LOW, MEDIUM, HIGH) rather than probability.")

    # Section 3: System Architecture
    h3 = doc.add_heading('3. System Architecture & Flow Design', level=1)
    h3.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph(
        "The microservice exposes REST endpoints using FastAPI. Downstream integrations query the API "
        "passing internal keys. Ingestion is handled asynchronously in background worker threads, storing "
        "documents as vectors in Qdrant."
    )
    doc.add_paragraph(
        "Embeddings are computed using NeuML/pubmedbert-base-embeddings (768 dimensions, cosine distance) "
        "for dense vectors and Qdrant FastEmbed BM25 for sparse indices. Merging is achieved via Reciprocal "
        "Rank Fusion (RRF)."
    )

    # Section 4: Document Ingestion Strategy
    h4 = doc.add_heading('4. Document Ingestion & Chunking Strategy', level=1)
    h4.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph(
        "To preserve detail and metadata, Guidelines are processed page-by-page. Repeated headers, footers, "
        "and page numbers are eliminated using frequency analysis before token boundaries are set."
    )
    doc.add_paragraph(
        "We enforce a 260-token target chunk with 45-token overlaps. The chunking algorithm preserves sections "
        "and prevents mixing child (pediatric WHO guidelines) and adult (ICMR guidelines) documents."
    )

    # Section 5: API Documentation
    h5 = doc.add_heading('5. API Reference & Operations', level=1)
    h5.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph(
        "Endpoints are secured using header validation. All protected business, management, and indexing APIs check "
        "X-Internal-API-Key. Upload APIs support multipart image/PDF streams and process them memory-only, ensuring patient privacy."
    )

    # Save
    doc.save(output_path)
    print(f"Success: Documentation compiled at {output_path}")

if __name__ == "__main__":
    main()
